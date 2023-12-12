from copy import deepcopy
from typing import List, Dict, Union
import re

import logging
import os
import log_config
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = log_config.setup_logger(__name__, logging.DEBUG)


class HandleDocument:
    """
    Class that handles the creation of the doc files, based on the CSV and JPG files
    """

    @staticmethod
    def add_csv_to_doc(filename: str, document: Document):
        """
        Add the CSV file contents to the document
        :param filename: The name of the CSV file
        :param document: The document object
        """
        # Read the CSV file into a DataFrame
        df = pd.read_csv(filename)
        # Add the DataFrame as a table to the document
        table = document.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        # Add header row
        for i, column_name in enumerate(df.columns):
            if i == 0:
                hdr_cells[i].text = 'Year'
            else:
                hdr_cells[i].text = str(column_name)
        # Add the rest of the data frame
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    @staticmethod
    def add_jpg_to_doc(filename, document):
        """
        Add the JPG file contents to the document
        :param filename: The name of the JPG file
        :param document: The document object
        """
        # Add the JPG file to the document
        document.add_picture(filename, width=Inches(8))  # You can adjust the width


class GenerateGraphs:
    CSV_DIR = "csvs"
    DOCS_DIR = "docs"
    YEARS = list(range(2012, 2023))
    DATAFRAME_COLUMNS = {
        "gender": ["Gender", "Gender/Woman Board", "Gender/Woman Executive"],
        "inclusivity": ["Inclusivity", "Inclusivity Board", "Inclusivity Executive"],
        "sustainability": ["Sustainability", "Sustainability Board", "Sustainability Executive"],
        "esg": ["ESG", "ESG Board", "ESG Executive"]
    }

    REGEX_PATTERNS = {
        "gender": [r"\bgender[a-z]*\b", r"\bwom[ae]n\b", r"\bboard(s)?\b", r"\bexecutiv[a-z]+"],
        "inclusivity": [r"\binclusiv[a-z]+\b", r"\bboard(s)?\b", r"\bexecutiv[a-z]+"],
        "sustainability": [r"\bsustain[a-z]+\b", r"\bboard(s)?\b", r"\bexecutiv[a-z]+"],
        "esg": [r"\besg\b", r"\bboard(s)?\b", r"\bexecutiv[a-z]+"]
    }

    def __init__(self, location: str):
        """
        Initialize the class
        :param location: The location where the pdfs have been decoded into txt
        """
        self.location = location
        self._companies = []
        files = os.listdir(self.location)
        self.files = [f for f in files if f.endswith("txt")]
        self.dfs = {}

    @property
    def companies(self) -> List:
        """
        Property that returns the list of companies that have been analysed
        :return: The list of individual companies
        """
        for f in self.files:
            company = next(iter(f.split("-")))
            if company not in self._companies:
                self._companies.append(company)
        return self._companies

    def extract_values(self, file: str, key: str) -> List:
        """
        Extract the needed values for the company
        :param file: The name of the file containing the PDF decoded into txt format
        :param key: The key for the dataframe determining the search terms
        :return: A list of values for the company
        """
        values = []
        for idx in range(3):
            if len(self.REGEX_PATTERNS[key]) == 4:
                main_pattern = re.compile(self.REGEX_PATTERNS[key][0], re.IGNORECASE)
                secondary_pattern = re.compile(self.REGEX_PATTERNS[key][1], re.IGNORECASE)
                board_pattern = re.compile(self.REGEX_PATTERNS[key][2], re.IGNORECASE)
                executive_pattern = re.compile(self.REGEX_PATTERNS[key][3], re.IGNORECASE)
            else:
                main_pattern = re.compile(self.REGEX_PATTERNS[key][0], re.IGNORECASE)
                secondary_pattern = None
                board_pattern = re.compile(self.REGEX_PATTERNS[key][1], re.IGNORECASE)
                executive_pattern = re.compile(self.REGEX_PATTERNS[key][2], re.IGNORECASE)

            with open(f"{self.location}/{file}", encoding="utf8") as f:
                text = f.read()
                value = 0
                for paragraph in text.split("\n"):
                    if idx == 0:
                        # first column, just a simple count for the main pattern
                        value += len(main_pattern.findall(paragraph))
                    elif idx == 1:
                        # second column, if secondary pattern is None, then just a simple count for the main + board pattern
                        # else, count the main + secondary + board pattern
                        if secondary_pattern is None:
                            if main_pattern.search(paragraph) and board_pattern.search(paragraph):
                                value += 1
                        else:
                            if (main_pattern.search(paragraph) or secondary_pattern.search(
                                    paragraph)) and board_pattern.search(paragraph):
                                value += 1
                    elif idx == 2:
                        # third column, if secondary pattern is None, then just a simple count for the main + executive pattern
                        # else, count the main + secondary + executive pattern
                        if secondary_pattern is None:
                            if main_pattern.search(paragraph) and executive_pattern.search(paragraph):
                                value += 1
                        else:
                            if (main_pattern.search(paragraph) or secondary_pattern.search(
                                    paragraph)) and executive_pattern.search(paragraph):
                                value += 1
                values.append(value)

        return values

    def update_df(self, file: Union[str, None], company: str, year: int, add_empty=False):
        """
        Updates the dataframe for that particular company with the results
        :param file: The name of the file containing the PDF decoded into txt format
        :param company: The name of the company file in the PDF file, eg REP for REPSOL
        :param year: The year as coded in the PDF file name
        :param add_empty: bool to add an empty row of Zeroes if there is no PDF for that company that year
        """
        for key in self.DATAFRAME_COLUMNS:
            df_id = f"{key}-{company}"
            if df_id not in self.dfs:
                self.dfs[df_id] = pd.DataFrame(columns=self.DATAFRAME_COLUMNS[key], index=self.YEARS)
            if add_empty:
                self.dfs[df_id].loc[year] = 0
                continue
            values = self.extract_values(file, key)
            self.dfs[df_id].loc[year] = values
        return

    def generate_csvs_for_company(self, company_name: str, force_generate: bool = False):
        """
        Generates the CSV files for the company, based on the self.REGEX_PATTERNS and saves them in the CSV_DIR
        :param company_name: The name of the company
        :param force_generate: bool, set to True to generate the CSV even is found
        :return:
        """
        company_csvs = os.listdir(self.CSV_DIR)
        company_csvs = [f for f in company_csvs if f.endswith("csv") and company_name in f]
        if not force_generate and len(company_csvs) == len(self.DATAFRAME_COLUMNS):
            logger.info(f"The CSV files for company {company_name} are already generated")
            return
        logger.info(f"Generating CSV files for company {company_name}")
        company_files = [f for f in self.files if company_name in f]
        for year in self.YEARS:
            for file in company_files:
                if str(year) in file:
                    # found the file for that particular year
                    logger.debug(f"Found data for company {company_name} for year {year}")
                    self.update_df(file, company_name, year)
                    break
            else:
                logger.debug(f"Did not find data for company {company_name} for year {year}")
                self.update_df(None, company_name, year, add_empty=True)

        for df_id in self.dfs:
            if company_name in df_id:
                self.dfs[df_id].to_csv(f"{self.CSV_DIR}/{df_id}.csv")

    def generate_graphs(self, company_name: str, force_generate: bool = False):
        """
        Generates the graph for the company
        :param company_name: The name of the company
        :param force_generate: bool, set to True to generate the CSV even is found
        """
        files = os.listdir(self.CSV_DIR)
        files = [f for f in files if f.endswith("csv") and company_name in f]

        for file in files:
            if os.path.exists(f"{self.CSV_DIR}/{file.split('.csv')[0]}_graph.jpg") and not force_generate:
                logger.info(f"The graph for company {company_name} based on CSV file {file} is already generated")
                continue
            logger.info(f"Generating graph for company {company_name} based on CSV file {file}")
            # Read the CSV file into a DataFrame
            df = pd.read_csv(f"{self.CSV_DIR}/{file}", index_col=0)
            # Plotting the graph
            plt.figure(figsize=(10, 6))  # You can change the figure size as needed
            # Loop through each column (except the index) and plot
            for column in df.columns:
                plt.plot(df.index, df[column], marker='o', label=column)
            # Adding labels and title
            plt.xlabel('Year')
            plt.ylabel('Values')
            plt.title(f'{company_name} Data Over Years')
            plt.legend()
            plt.grid(True)
            # Save the figure
            plt.savefig(f"{self.CSV_DIR}/{file.split('.csv')[0]}_graph.jpg")
            # Close the figure to prevent the warning
            plt.close()

    def generate_doc(self, company_name: str, force_generate: bool = False):
        """"
        Generates the doc files for the company that includes the graphs and the csv data into a single doc
        :param company_name: The name of the company
        :param force_generate: bool, set to True to generate the CSV even is found
        """
        logger.info(f"Generating the doc file for company {company_name}")
        if not os.path.exists(f"{self.DOCS_DIR}"):
            logger.error(f"The directory {self.DOCS_DIR} does not exist, creating it")
            os.mkdir(self.DOCS_DIR)

        if not force_generate and os.path.exists(f"{self.DOCS_DIR}/{company_name}.docx"):
            logger.info(f"The doc file for company {company_name} is already generated")
            return
        files = os.listdir(self.CSV_DIR)
        csv_files = [f for f in files if f.endswith("csv") and company_name in f]
        jpg_files = [f for f in files if f.endswith("jpg") and company_name in f]
        csv_files.sort()
        jpg_files.sort()

        # Create a new Document
        doc = Document()
        title = doc.add_heading(f'{company_name} Results and Graphs', 0)
        for run in title.runs:
            run.font.bold = True
            run.font.size = Pt(26)  # Optional: Adjust the font size if needed
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")

        paragraph = doc.add_paragraph(
            f'The following graphs and data are based on collected data from the PDF files for the company: '
            f'{company_name.upper()}')
        paragraph.style.font.size = Pt(14)

        doc.add_paragraph("")

        for csv_file, jpg_file in zip(csv_files, jpg_files):
            # Add CSV and JPG files to the document
            pattern = csv_file.split("-")[0].upper()
            doc.add_heading(f'{pattern} COLLECTED DATA TABLE:', level=1)
            HandleDocument.add_csv_to_doc(f"{self.CSV_DIR}/{csv_file}", doc)
            doc.add_heading(f'{pattern} GRAPH:', level=1)
            HandleDocument.add_jpg_to_doc(f"{self.CSV_DIR}/{jpg_file}", doc)
            doc.add_page_break()  # Add a page break after each CSV/jpg content

        # Save the document
        doc.save(f'{self.DOCS_DIR}/{company_name}.docx')
        logger.info(f'The Word document has been created with the CSV and JPG files for {company_name}.')

    def analyse_and_plot_data_for_company(self, company_name: str, force_generate: bool = False):
        """
        Main function in the class, that does the following:
        1. Gets the files and searches for the key terms
        2. Puts the results into CSV files
        3. Create graphs based on the CSV data
        4. Populate doc files with these graphs
        :param company_name: String identifing the company, same as in the PDF file, REP for Reposol
        :param force_generate: bool, set to True to generate the CSV even is found
        """
        logger.info(f"Analysing and plotting data for company {company_name}")
        self.generate_csvs_for_company(company_name, force_generate)
        self.generate_graphs(company_name, force_generate)
        self.generate_doc(company_name, force_generate)

    def generate_aggregated_doc(self):
        """
        Go over all the companies and generate an aggregated doc file with a table and graph adding the existing csvs
        """
        logger.info(f"Generating the aggregated doc file")
        # create csv files for each of the dataframes
        for key in self.DATAFRAME_COLUMNS:
            df_id = f"{key}-all"
            if df_id not in self.dfs:
                self.dfs[df_id] = pd.DataFrame(columns=self.DATAFRAME_COLUMNS[key], index=self.YEARS)
                # set it all to zeroes
                for year in self.YEARS:
                    self.dfs[df_id].loc[year] = 0
            for company in self.companies:
                company_df = pd.read_csv(f"{self.CSV_DIR}/{key}-{company}.csv", index_col=0)
                # sum the values for each year
                for year in self.YEARS:
                    self.dfs[df_id].loc[year] += company_df.loc[year]
            # generate the csv file
            self.dfs[df_id].to_csv(f"{self.CSV_DIR}/{df_id}.csv")
        self.generate_graphs("all")
        self.generate_doc("all")


if __name__ == '__main__':

    # testing out the creation of the doc file
    doc = Document()

    csv_files = os.listdir("csvs")
    csv_files = [f"csvs/{f}" for f in csv_files if f.endswith("csv")]
    jpg_files = os.listdir("csvs")
    jpg_files = [f"csvs/{f}" for f in jpg_files if f.endswith("jpg")]

    # Add CSV and JPG files to the document
    for csv_file in csv_files:
        HandleDocument.add_csv_to_doc(csv_file, doc)
        doc.add_page_break()  # Add a page break after each CSV content

    for jpg_file in jpg_files:
        HandleDocument.add_jpg_to_doc(jpg_file, doc)
        doc.add_page_break()  # Add a page break after each image

    # Save the document
    doc.save('Report.docx')

    print('The Word document has been created with the CSV and JPG files.')
